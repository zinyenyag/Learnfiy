from fastapi import FastAPI, File, UploadFile, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from PyPDF2 import PdfReader
import base64
import re
import os
from urllib.parse import quote
import requests
from dotenv import load_dotenv
import json
from typing import List, Optional
from datetime import datetime
import tempfile
import io
import time

# Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Load environment variables
load_dotenv()

# AI Provider imports (with graceful fallback)
try:
    from openai import OpenAI  # type: ignore[reportMissingImports]
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    import google.generativeai as genai  # type: ignore[reportMissingImports]
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

try:
    from anthropic import Anthropic  # type: ignore[reportMissingImports]
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

app = FastAPI()

# Allow your index.html (localhost) to call the backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # tighten later
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

SYLLABUS_PDFS = {
    "Mathematics": "Mathematics Syllabus.pdf",
    "Accounting": "Accounting Syllabus.pdf",
    "Economics": "Economics Syllabus.pdf",
    "Computer Science": "Computer Science Syllabus.pdf",
    "Business Studies": "Business Studies Syllabus.pdf"
}

_cached_chunks = {}  # Cache per subject

# Initialize AI clients
openai_client = None
gemini_client = None
anthropic_client = None

if OPENAI_AVAILABLE and os.getenv("OPENAI_API_KEY"):
    openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

if GEMINI_AVAILABLE and os.getenv("GEMINI_API_KEY"):
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
    gemini_client = genai.GenerativeModel('gemini-pro')

if ANTHROPIC_AVAILABLE and os.getenv("ANTHROPIC_API_KEY"):
    anthropic_api_key = os.getenv("ANTHROPIC_API_KEY")
    if anthropic_api_key and anthropic_api_key != "sk-ant-PASTE_YOUR_KEY_HERE":
        anthropic_client = Anthropic(api_key=anthropic_api_key)

def extract_text_from_pdf(pdf_path: str) -> str:
    reader = PdfReader(pdf_path)
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text())
    return "\n".join(parts)

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a .docx file using python-docx"""
    if not DOCX_AVAILABLE:
        return "Error: python-docx library not available. Cannot read .docx files."
    try:
        doc = Document(docx_path)
        parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n".join(parts)
    except Exception as e:
        return f"Error extracting text from .docx file: {str(e)}"

def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 200):
    text = re.sub(r"\n{3,}", "\n\n", text)
    chunks = []
    i = 0
    while i < len(text):
        chunk = text[i:i+chunk_size]
        chunks.append(chunk.strip())
        i += (chunk_size - overlap)
    return [c for c in chunks if len(c) > 100]

def score_chunk(query: str, chunk: str) -> float:
    # Simple keyword scoring (fast baseline). We'll upgrade to embeddings later.
    q = re.findall(r"[a-zA-Z0-9]+", query.lower())
    c = chunk.lower()
    if not q:
        return 0.0
    hits = sum(1 for w in set(q) if w in c)
    return hits / max(6, len(set(q)))  # normalized

def get_chunks(subject: str = "Mathematics"):
    global _cached_chunks
    if subject not in _cached_chunks:
        pdf_name = SYLLABUS_PDFS.get(subject, "Mathematics Syllabus.pdf")
        # PDFs are in parent directory (project root)
        pdf_path = os.path.join(os.path.dirname(__file__), "..", pdf_name)
        if os.path.exists(pdf_path):
            try:
                text = extract_text_from_pdf(pdf_path)
                _cached_chunks[subject] = chunk_text(text)
            except Exception as e:
                print(f"Error loading PDF for {subject}: {e}")
                _cached_chunks[subject] = []
        else:
            _cached_chunks[subject] = []
    return _cached_chunks.get(subject, [])

def retrieve_top_k(query: str, subject: str = "Mathematics", k: int = 4):
    chunks = get_chunks(subject)
    if not chunks:
        return []
    scored = [(score_chunk(query, ch), ch) for ch in chunks]
    scored.sort(key=lambda x: x[0], reverse=True)
    return scored[:k]

def Web_Search_duckduckgo(query: str, k: int = 5):
    # No API key needed. Returns limited results.
    # If it fails, we just return empty list.
    try:
        url = f"https://duckduckgo.com/html/?q={quote(query)}"
        r = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
        html = r.text
        # Very light parsing:
        links = re.findall(r'class="result__a" href="(.*?)"', html)
        cleaned = []
        for link in links:
            if link.startswith("http"):
                cleaned.append(link)
            if len(cleaned) >= k:
                break
        return cleaned
    except Exception:
        return []

def call_ai_openai(prompt: str, context: str = "") -> str:
    """Call OpenAI ChatGPT API"""
    if not openai_client:
        return None
    
    try:
        system_prompt = """You are an expert tutor for Forms 5-6 students in Zimbabwe. Answer questions directly and completely. Provide full solutions with all steps and final answers."""
        
        # Use the prompt directly - it's already formatted with context
        full_prompt = prompt
        
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",  # Cost-effective, can change to gpt-4 for better quality
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": full_prompt}
            ],
            temperature=0.7,
            max_tokens=2000  # Increased for comprehensive solutions
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"OpenAI error: {e}")
        return None

def call_ai_gemini(prompt: str, context: str = "") -> str:
    """Call Google Gemini API"""
    if not gemini_client:
        return None
    
    try:
        # Use the prompt directly - it's already formatted
        full_prompt = prompt
        
        response = gemini_client.generate_content(
            f"Answer directly and completely. {full_prompt}"
        )
        return response.text
    except Exception as e:
        print(f"Gemini error: {e}")
        return None

def call_ai_anthropic(prompt: str, context: str = "", subject: str = "Mathematics") -> str:
    """Call Anthropic Claude API using the latest format"""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key or api_key == "sk-ant-PASTE_YOUR_KEY_HERE":
        print("Anthropic: API key not found or is placeholder")
        return None
    
    try:
        client = Anthropic(api_key=api_key)
        
        # Create system prompt
        system_prompt = f"You are a helpful tutor for Forms 5-6 {subject} students in Zimbabwe. Answer questions directly, clearly, and completely. Provide full solutions with all steps and final answers."
        
        # Use the prompt directly - it's already formatted
        user_prompt = prompt
        
        # Try the latest model first, fallback to haiku if not available
        try:
            message = client.messages.create(
                model="claude-3-5-sonnet-20241022",  # Latest Claude Sonnet model
                max_tokens=2000,
                temperature=0.3,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}]
            )
            print("Anthropic: Used claude-3-5-sonnet-20241022")
        except Exception as e1:
            print(f"Anthropic: Sonnet model failed, trying haiku: {e1}")
            # Fallback to haiku if sonnet not available
            message = client.messages.create(
                model="claude-3-haiku-20240307",  # Fast and cost-effective
                max_tokens=2000,
                temperature=0.3,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}]
            )
            print("Anthropic: Used claude-3-haiku-20240307")
        
        # Extract text from message content blocks
        parts = []
        for block in message.content:
            if getattr(block, "type", None) == "text":
                parts.append(block.text)
        result = "\n".join(parts).strip()
        print(f"Anthropic: Success, response length: {len(result)}")
        return result
    except Exception as e:
        print(f"Anthropic error: {e}")
        import traceback
        traceback.print_exc()
        return None

def call_ai_with_fallback(prompt: str, context: str = "", subject: str = "Mathematics") -> str:
    """
    Auto-mode: Try AI providers in order of preference with fallback
    Priority: OpenAI > Gemini > Anthropic > Fallback response
    """
    # Try OpenAI first
    if openai_client:
        result = call_ai_openai(prompt, context)
        if result:
            return result
    
    # Fallback to Gemini
    if gemini_client:
        result = call_ai_gemini(prompt, context)
        if result:
            return result
    
    # Fallback to Anthropic
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if api_key and api_key != "sk-ant-PASTE_YOUR_KEY_HERE":
        result = call_ai_anthropic(prompt, context, subject)
        if result:
            return result
    
    # If no AI available, provide intelligent fallback
    return generate_fallback_response(prompt, context, subject)

def get_subject_formatting_instructions(subject: str) -> str:
    """Get subject-specific formatting instructions for AI responses"""
    instructions = {
        "Accounting": """**ACCOUNTING FORMATTING (IAS 1 Standards)**:
- Present financial statements in proper columnar format with clear labels
- Use proper headings: "Income Statement", "Statement of Financial Position", etc.
- Format numbers with proper alignment (right-align numbers, left-align descriptions)
- Use clear sections: Trading Account, Profit and Loss Account, Balance Sheet
- Show calculations clearly with proper line items
- Use proper accounting terminology and presentation standards
- Format as:
  ```
  INCOME STATEMENT
  For the year ended [Date]
  
  Sales                          $XXX
  Less: Cost of Sales           ($XXX)
  Gross Profit                  $XXX
  
  Less: Expenses
    Depreciation                 $XXX
    Other expenses               $XXX
  Net Profit                    $XXX
  ```
- Make it look like a professional accounting document that can be printed or exported to Word""",
        
        "Mathematics": """**MATHEMATICS FORMATTING**:
- Show EVERY step of the solution clearly
- Number each step: Step 1, Step 2, Step 3, etc.
- After each step, provide a brief explanation of what was done and why
- Use proper mathematical notation and formatting
- Show all calculations, substitutions, and simplifications
- Format as:
  ```
  Question: [Question text]
  
  Solution:
  
  Step 1: [Action taken]
  Explanation: [Why this step was taken]
  
  Step 2: [Action taken]
  Explanation: [Why this step was taken]
  
  Step 3: [Action taken]
  Explanation: [Why this step was taken]
  
  Final Answer: [Answer]
  ```
- Make it clear, educational, and easy to follow""",
        
        "Economics": """**ECONOMICS FORMATTING**:
- Use professional academic writing style
- Structure with clear headings and subheadings
- Use proper economic terminology and concepts
- Include diagrams descriptions where relevant (explain what the diagram would show)
- Format definitions clearly
- Use bullet points for lists of factors, causes, effects
- Format as:
  ```
  Topic: [Topic Name]
  
  Definition:
  [Clear definition]
  
  Key Concepts:
  • Concept 1: [Explanation]
  • Concept 2: [Explanation]
  
  Analysis:
  [Detailed analysis following syllabus requirements]
  
  Conclusion:
  [Summary]
  ```
- Make it look like professional academic notes""",
        
        "Business Studies": """**BUSINESS STUDIES FORMATTING**:
- Use professional business writing style
- Structure with clear headings: Introduction, Analysis, Conclusion
- Use business terminology appropriately
- Format case studies clearly
- Use bullet points for lists (SWOT, PEST, etc.)
- Format as:
  ```
  Topic: [Topic Name]
  
  Introduction:
  [Context and overview]
  
  Main Points:
  1. [Point 1 with explanation]
  2. [Point 2 with explanation]
  
  Analysis:
  [Detailed analysis]
  
  Conclusion:
  [Summary and recommendations]
  ```
- Make it look like professional business documentation""",
        
        "Computer Science": """**COMPUTER SCIENCE FORMATTING**:
- Use proper code formatting with clear indentation
- Explain code line by line where relevant
- Use proper technical terminology
- Format algorithms clearly
- Format as:
  ```
  Problem: [Problem description]
  
  Solution:
  
  Algorithm:
  1. [Step 1]
  2. [Step 2]
  3. [Step 3]
  
  Code:
  ```[language]
  [Properly formatted code]
  ```
  
  Explanation:
  - Line 1: [Explanation]
  - Line 2: [Explanation]
  
  Time Complexity: O(...)
  Space Complexity: O(...)
  ```
- Make it look like professional technical documentation"""
    }
    
    return instructions.get(subject, """**GENERAL FORMATTING**:
- Use clear headings and subheadings
- Structure content logically
- Use bullet points for lists
- Format numbers and calculations clearly
- Make it professional and easy to read""")

def generate_fallback_response(prompt: str, context: str, subject: str) -> str:
    """Generate a helpful response when AI APIs are not available"""
    # Provide a direct answer based on context if available
    if context:
        return f"""Based on the {subject} syllabus:

{context[:800]}...

Note: For more detailed answers, please set up an AI API key (OpenAI, Gemini, or Anthropic) in your .env file."""
    
    return f"""I need an AI API key to provide detailed answers. Please set up OPENAI_API_KEY, GEMINI_API_KEY, or ANTHROPIC_API_KEY in your .env file.

For {subject} questions, I can help once the API is configured."""

def format_answer_for_subject(answer: str, subject: str) -> str:
    """Format the AI answer with subject-specific HTML formatting"""
    # Convert markdown-style formatting to HTML
    formatted = answer
    
    # Convert code blocks
    formatted = re.sub(r'```(\w+)?\n(.*?)```', r'<pre style="background: #f5f5f5; padding: 15px; border-radius: 5px; overflow-x: auto;"><code>\2</code></pre>', formatted, flags=re.DOTALL)
    
    # Convert bold
    formatted = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', formatted)
    
    # Convert headings
    formatted = re.sub(r'^### (.*?)$', r'<h3 style="color: #667eea; margin-top: 20px;">\1</h3>', formatted, flags=re.MULTILINE)
    formatted = re.sub(r'^## (.*?)$', r'<h2 style="color: #667eea; margin-top: 20px; border-bottom: 2px solid #667eea; padding-bottom: 5px;">\1</h2>', formatted, flags=re.MULTILINE)
    formatted = re.sub(r'^# (.*?)$', r'<h1 style="color: #667eea; margin-top: 20px;">\1</h1>', formatted, flags=re.MULTILINE)
    
    # Convert line breaks
    formatted = formatted.replace('\n\n', '</p><p style="margin: 10px 0; line-height: 1.6;">')
    formatted = f'<p style="margin: 10px 0; line-height: 1.6;">{formatted}</p>'
    
    # Subject-specific formatting
    if subject == "Accounting":
        # Format financial statements with proper spacing
        formatted = re.sub(r'(\$[\d,]+\.?\d*)', r'<span style="text-align: right; font-family: monospace; font-weight: bold;">\1</span>', formatted)
        formatted = formatted.replace('INCOME STATEMENT', '<h2 style="text-align: center; font-weight: bold; color: #333; margin: 20px 0;">INCOME STATEMENT</h2>')
        formatted = formatted.replace('Statement of Financial Position', '<h2 style="text-align: center; font-weight: bold; color: #333; margin: 20px 0;">Statement of Financial Position</h2>')
        # Format columnar data
        formatted = re.sub(r'^([A-Z][^:]+):\s+(\$?[\d,]+\.?\d*)$', r'<div style="display: flex; justify-content: space-between; margin: 5px 0;"><span>\1</span><span style="text-align: right; font-weight: bold;">\2</span></div>', formatted, flags=re.MULTILINE)
    
    elif subject == "Mathematics":
        # Format steps clearly
        formatted = re.sub(r'Step (\d+):', r'<div style="margin: 15px 0; padding: 12px; background: linear-gradient(135deg, #f0f4ff 0%, #e0e7ff 100%); border-left: 4px solid #667eea; border-radius: 5px;"><strong style="color: #667eea;">Step \1:</strong>', formatted)
        formatted = re.sub(r'Explanation:', r'</div><div style="margin: 8px 0 20px 25px; padding: 8px; background: #f9fafb; border-left: 3px solid #22c55e; color: #666; font-style: italic;">Explanation:', formatted)
        formatted = re.sub(r'Final Answer:', r'</div><div style="margin: 15px 0; padding: 12px; background: linear-gradient(135deg, #dcfce7 0%, #bbf7d0 100%); border: 2px solid #22c55e; border-radius: 5px; font-weight: bold; color: #166534;"><strong>Final Answer:</strong>', formatted)
    
    return formatted

class ChatRequest(BaseModel):
    subject: str = "Mathematics"
    message: str

class ChatResponse(BaseModel):
    answer: str
    used_syllabus: bool
    syllabus_snippets: list[str]
    external_sources: list[str]

@app.post("/api/chat", response_model=ChatResponse)
async def chat(request: Request):
    """Chat endpoint that accepts text messages and optional file uploads.
    Supports both JSON (no files) and FormData (with files) requests."""
    
    try:
        content_type = request.headers.get("content-type", "")
        file_contents = []
        transcribed_audio_text = ""
        query = ""
        subject = "Mathematics"
        form = None  # Initialize form to avoid scope issues
        audio_file = None  # Initialize audio_file to avoid scope issues
        
        # Check if it's FormData (multipart/form-data)
        files = []
        if "multipart/form-data" in content_type:
            try:
                form = await request.form()
            except Exception as form_error:
                print(f"Error parsing form data: {form_error}")
                form = None
            
            if form is not None:
                query = (form.get("message") or "").strip()
                subject = (form.get("subject") or "Mathematics")
                
                # Process uploaded files
                files = form.getlist("files")
                
                # Process audio file if present
                audio_file = form.get("audio_file")
        
        # Process audio file if present (outside form check to handle scope)
        if audio_file and openai_client:
                try:
                    import tempfile
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.webm') as tmp_audio:
                        tmp_audio.write(await audio_file.read())
                        tmp_audio_path = tmp_audio.name
                    
                    with open(tmp_audio_path, "rb") as audio_f:
                        transcription = openai_client.audio.transcriptions.create(
                            model="whisper-1", 
                            file=audio_f
                        )
                        transcribed_audio_text = transcription.text
                        print(f"Transcribed Audio: {transcribed_audio_text}")
                except Exception as e:
                    print(f"Error transcribing audio: {e}")
                    transcribed_audio_text = f"Error transcribing audio: {str(e)}"
                finally:
                    if 'tmp_audio_path' in locals() and os.path.exists(tmp_audio_path):
                        try:
                            os.unlink(tmp_audio_path)
                        except:
                            pass
        
        # Process uploaded files (if any)
        if files:
            for file in files:
                try:
                    content = await file.read()
                    file_name = file.filename
                    
                    # Extract text based on file type
                    if file_name.endswith('.pdf'):
                        import tempfile
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                            tmp.write(content)
                            tmp_path = tmp.name
                        try:
                            text = extract_text_from_pdf(tmp_path)
                            # Extract more content for better analysis (up to 10000 chars for comprehensive analysis)
                            if len(text) > 10000:
                                # Get beginning and end sections for very long documents
                                extracted_text = text[:7000] + "\n\n... [MIDDLE CONTENT OMITTED] ...\n\n" + text[-3000:]
                                extracted_text += f"\n\n(Note: Full document has {len(text)} characters. Showing key sections.)"
                            else:
                                extracted_text = text
                            file_contents.append(f"Document: {file_name}\n\nFull Content:\n{extracted_text}")
                        except Exception as e:
                            file_contents.append(f"File {file_name} uploaded. Error extracting text: {str(e)}. Please analyze the PDF file.")
                        finally:
                            if os.path.exists(tmp_path):
                                try:
                                    os.unlink(tmp_path)
                                except:
                                    pass
                    elif file_name.endswith('.docx'):
                        # Handle .docx files using python-docx
                        import tempfile
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                            tmp.write(content)
                            tmp_path = tmp.name
                        try:
                            text = extract_text_from_docx(tmp_path)
                            # Extract more content for better analysis (up to 10000 chars)
                            if len(text) > 10000:
                                # Get beginning and end sections for very long documents
                                extracted_text = text[:7000] + "\n\n... [MIDDLE CONTENT OMITTED] ...\n\n" + text[-3000:]
                                extracted_text += f"\n\n(Note: Full document has {len(text)} characters. Showing key sections.)"
                            else:
                                extracted_text = text
                            file_contents.append(f"Document: {file_name}\n\nFull Content:\n{extracted_text}")
                        except Exception as e:
                            file_contents.append(f"File {file_name} uploaded. Error extracting text: {str(e)}. Please provide complete solutions to all problems in this Word document.")
                        finally:
                            if os.path.exists(tmp_path):
                                try:
                                    os.unlink(tmp_path)
                                except:
                                    pass
                    elif file_name.endswith(('.txt', '.doc')):
                        # Handle plain text and old .doc files
                        try:
                            text = content.decode('utf-8')
                            # Get more content for text files too
                            extracted_text = text[:5000] if len(text) > 5000 else text
                            file_contents.append(f"Content from {file_name}:\n{extracted_text}")
                        except:
                            try:
                                # Try different encodings
                                text = content.decode('latin-1')
                                extracted_text = text[:5000] if len(text) > 5000 else text
                                file_contents.append(f"Content from {file_name}:\n{extracted_text}")
                            except:
                                file_contents.append(f"File {file_name} uploaded but could not read as text. Please analyze this file.")
                    elif file_name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                        file_contents.append(f"Image file {file_name} uploaded. Please provide complete solutions to all problems, questions, or exercises shown in this image. Solve each problem fully with step-by-step solutions.")
                    elif file_name.lower().endswith(('.webm', '.mp3', '.wav', '.ogg', '.m4a')):
                        # Audio file - note that transcription would need additional processing
                        file_contents.append(f"Audio/voice message file {file_name} uploaded. Please transcribe and provide complete solutions to any questions or problems asked in the audio.")
                    else:
                        file_contents.append(f"File {file_name} uploaded. Please provide complete, detailed solutions to all problems, questions, or exercises contained in this file. Solve each one fully.")
                except Exception as e:
                    file_contents.append(f"Error processing {file.filename}: {str(e)}. Please analyze the uploaded file.")
        
        # Handle JSON request (text-only) if not multipart/form-data
        if "multipart/form-data" not in content_type:
            try:
                body = await request.json()
                query = body.get("message", "").strip()
                subject = body.get("subject", "Mathematics")
            except:
                # Fallback if JSON parsing fails
                query = ""
                subject = "Mathematics"
        
        # Combine query, transcribed audio, and file content
        full_query_parts = []
        if query:
            full_query_parts.append(query)
        if transcribed_audio_text:
            full_query_parts.append(f"Transcribed voice message: {transcribed_audio_text}")
        if file_contents:
            full_query_parts.append("Attached files:\n" + "\n\n".join(file_contents))

        full_query = "\n\n".join(full_query_parts) if full_query_parts else "Please help me learn."
        
        if not full_query.strip():
            return ChatResponse(
                answer="Please provide a question or attach a file to analyze.",
                used_syllabus=False,
                syllabus_snippets=[],
                external_sources=[]
            )

        # Retrieve relevant syllabus content
        top = retrieve_top_k(full_query, subject=subject, k=4)
        best_score = top[0][0] if top else 0.0
        snippets = [ch[:600] for s, ch in top if s > 0]  # short preview
        
        # Build context from syllabus
        context = ""
        if snippets:
            context = "\n\n".join(snippets[:3])  # Use top 3 snippets as context
        
        used_syllabus = best_score >= 0.18 or len(snippets) > 0

        # Build subject-specific formatting instructions
        formatting_instructions = get_subject_formatting_instructions(subject)
        
        # Build the prompt with subject context and file content
        user_request = query if query else "Please provide complete solutions to all problems, questions, or exercises in the uploaded document."
        if transcribed_audio_text:
            user_request += f"\n\n(User also said via voice: {transcribed_audio_text})"

        # Always get web search results for comprehensive answers
        web_sources = Web_Search_duckduckgo(f"{subject} {query}", k=5) if query else []
        web_context = ""
        if web_sources:
            web_context = f"\n\n**Additional Web Resources Available**: {', '.join(web_sources[:3])}"

        # Build a clean, direct prompt that won't be echoed
        if file_contents:
            # Document uploaded - provide complete solutions
            document_prompt = f"""Answer the following question and solve all problems in the uploaded document(s).

Question: {user_request}

Document Content:
{chr(10).join(file_contents)}

Syllabus Reference ({subject}):
{context if context else ""}

Provide complete solutions with all steps and final answers. Format according to {subject} standards."""
        else:
            # Direct question - answer directly
            document_prompt = f"""Answer this question directly and completely:

{user_request}

Syllabus Reference ({subject}):
{context if context else ""}

Provide a clear, complete answer with examples if needed."""
        
        enhanced_prompt = document_prompt
        
        # Call AI with auto-fallback (OpenAI -> Gemini -> Anthropic -> Fallback)
        # Context is already included in enhanced_prompt, so pass empty string
        answer = call_ai_with_fallback(enhanced_prompt, "", subject)
        
        # Return plain text answer (like ChatGPT) - no HTML formatting
        # The frontend will handle basic formatting for display
        
        # Always get external sources for comprehensive answers
        sources = Web_Search_duckduckgo(f"{subject} {query}", k=5) if query else []
        # Also search for document-related queries
        if file_contents and not sources:
            sources = Web_Search_duckduckgo(f"{subject} {query} solutions", k=5)

        return ChatResponse(
            answer=answer,  # Plain text, no HTML formatting
            used_syllabus=used_syllabus,
            syllabus_snippets=snippets,
            external_sources=sources
        )
    except Exception as e:
        import traceback
        error_msg = f"Error processing request: {str(e)}\n\nTraceback:\n{traceback.format_exc()}"
        print(error_msg)  # Log to console for debugging
        return ChatResponse(
            answer=f"I encountered an error while processing your request: {str(e)}\n\nPlease try again, or ensure your file is in a supported format (PDF, Word, Text, or Image). If the problem persists, check that the backend server is running correctly.",
            used_syllabus=False,
            syllabus_snippets=[],
            external_sources=[]
        )

class HomeworkRequest(BaseModel):
    subject: str = "Mathematics"
    question_text: str = ""
    question_type: str = "text"  # "text" or "file"
    file_name: str = ""  # Optional: name of uploaded file

class HomeworkResponse(BaseModel):
    analysis: str
    answer: str
    explanation: str
    score: str = ""
    mistakes: list[str] = []
    suggestions: list[str] = []

@app.post("/api/homework/analyze", response_model=HomeworkResponse)
async def analyze_homework(req: HomeworkRequest):
    """Analyze homework question and provide AI-generated answer"""
    subject = req.subject or "Mathematics"
    question = req.question_text.strip()
    
    # Handle file uploads - extract text if provided
    if req.question_type == "file" and req.file_name:
        if req.file_name.endswith('.pdf'):
            try:
                # If file content is provided as base64, decode and extract
                # For now, use the question_text which should contain extracted text
                if not question:
                    question = f"Analyze questions from uploaded file: {req.file_name}"
            except Exception as e:
                question = f"Please analyze the questions in the uploaded file: {req.file_name}. {question}"
    
    if not question:
        return HomeworkResponse(
            analysis="No question provided",
            answer="Please provide a question to analyze. You can type it in the text box or upload a file.",
            explanation="",
            mistakes=[],
            suggestions=[]
        )
    
    # Retrieve relevant syllabus content
    top = retrieve_top_k(question, subject=subject, k=4)
    snippets = [ch[:600] for s, ch in top if s > 0]
    context = "\n\n".join(snippets[:3]) if snippets else ""
    
    # Create comprehensive, subject-aware prompt for homework analysis
    subject_guidance = {
        "Mathematics": "Focus on mathematical concepts, formulas, step-by-step calculations, and problem-solving methods.",
        "Accounting": "Focus on accounting principles, double-entry bookkeeping, financial statements, and accounting standards (IAS).",
        "Economics": "Focus on economic theories, demand and supply, market structures, and macroeconomic concepts.",
        "Computer Science": "Focus on programming concepts, algorithms, data structures, and computational thinking.",
        "Business Studies": "Focus on business management, marketing strategies, organizational behavior, and business planning."
    }
    
    guidance = subject_guidance.get(subject, "Focus on the key concepts and principles relevant to this subject.")
    
    analysis_prompt = f"""As an expert tutor for Forms 5-6 {subject} students in Zimbabwe, analyze this homework question and provide MULTIPLE comprehensive solutions and approaches.

**Your Role:**
You are teaching {subject} at Forms 5-6 level. Generate AS MANY DIFFERENT SOLUTIONS as possible.

**IMPORTANT: Generate Multiple Solutions**
Provide at least 3-5 different approaches/solutions to this question. Each solution should:
- Use a different method or approach
- Be complete and step-by-step
- Explain why this particular approach works
- Be suitable for Forms 5-6 level

**Your Response Should Include:**

1. **Multiple Solutions** (at least 3-5 different approaches):
   - Solution 1: [Method/Approach 1]
   - Solution 2: [Method/Approach 2]
   - Solution 3: [Method/Approach 3]
   - Solution 4: [Alternative method if applicable]
   - Solution 5: [Another perspective/approach]

2. **Comparison of Methods**: Explain when to use each approach and their advantages

3. **Explanation**: For each solution, explain the underlying concepts

4. **Common Mistakes**: List what students typically get wrong

5. **Suggestions**: Tips for choosing the best approach

**Subject-Specific Focus:**
{guidance}

**Subject:** {subject}
**Question/Problem:** {question}

**Generate as many diverse solutions as possible. Be creative and show different ways to solve this problem!**"""

    # Get AI response with syllabus context
    ai_response = call_ai_with_fallback(analysis_prompt, context, subject)
    
    # Parse and structure the response
    answer_text = ai_response
    explanation = ai_response  # AI response includes both answer and explanation
    
    # Extract mistakes and suggestions from AI response if formatted
    mistakes = []
    suggestions = []
    
    # Try to parse structured response
    if "**Common Mistakes**" in answer_text or "Common Mistakes:" in answer_text:
        try:
            mistakes_section = answer_text.split("**Common Mistakes**")[1].split("**")[0] if "**Common Mistakes**" in answer_text else ""
            mistakes = [m.strip() for m in mistakes_section.split("\n") if m.strip() and m.strip().startswith(("-", "•", "*"))]
        except:
            pass
    
    if "**Suggestions**" in answer_text or "Suggestions:" in answer_text:
        try:
            suggestions_section = answer_text.split("**Suggestions**")[1].split("**")[0] if "**Suggestions**" in answer_text else ""
            suggestions = [s.strip() for s in suggestions_section.split("\n") if s.strip() and s.strip().startswith(("-", "•", "*"))]
        except:
            pass
    
    return HomeworkResponse(
        analysis=f"AI Analysis for {subject}",
        answer=answer_text,
        explanation=explanation,
        score="Analysis Complete ✓",
        mistakes=mistakes[:5],  # Limit to top 5
        suggestions=suggestions[:5]  # Limit to top 5
    )

class HomeworkAssignmentRequest(BaseModel):
    subject: str
    weak_areas: List[str] = []
    num_questions: int = 5
    difficulty: str = "medium"

@app.post("/api/homework/submit")
async def submit_homework(request: Request):
    """Submit homework - accepts and stores homework files. No automatic marking."""
    try:
        content_type = request.headers.get("content-type", "")
        
        if "multipart/form-data" not in content_type:
            return {"error": "Request must be multipart/form-data"}
        
        form = await request.form()
        subject = form.get("subject", "Mathematics")
        files = form.getlist("files")
        
        if not files or len(files) == 0:
            return {"error": "No files uploaded"}
        
        # Process files - extract text and compress images
        file_contents = []
        image_files = []
        text_content = ""
        
        for file in files:
            content = await file.read()
            file_name = file.filename
            
            if file_name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                # Store image for compression
                image_files.append({
                    'name': file_name,
                    'content': content
                })
                file_contents.append(f"Image file: {file_name} (will be analyzed)")
            elif file_name.endswith('.pdf'):
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                try:
                    text = extract_text_from_pdf(tmp_path)
                    text_content += f"\n\n--- Content from {file_name} ---\n{text[:10000]}"
                    file_contents.append(f"PDF: {file_name}")
                except Exception as e:
                    file_contents.append(f"PDF {file_name}: Error extracting text - {str(e)}")
                finally:
                    if os.path.exists(tmp_path):
                        try:
                            os.unlink(tmp_path)
                        except:
                            pass
            elif file_name.endswith('.docx'):
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                try:
                    text = extract_text_from_docx(tmp_path)
                    text_content += f"\n\n--- Content from {file_name} ---\n{text[:10000]}"
                    file_contents.append(f"Word: {file_name}")
                except Exception as e:
                    file_contents.append(f"Word {file_name}: Error extracting text - {str(e)}")
                finally:
                    if os.path.exists(tmp_path):
                        try:
                            os.unlink(tmp_path)
                        except:
                            pass
            elif file_name.endswith('.txt'):
                try:
                    text = content.decode('utf-8')
                    text_content += f"\n\n--- Content from {file_name} ---\n{text[:5000]}"
                    file_contents.append(f"Text: {file_name}")
                except:
                    file_contents.append(f"Text {file_name}: Could not decode")
        
        # Generate homework ID
        homework_id = f"hw_{subject.lower().replace(' ', '_')}_{int(time.time())}"
        
        # Return success response - homework uploaded, no marking
        return {
            "homework_id": homework_id,
            "subject": subject,
            "status": "uploaded",
            "message": f"Homework uploaded successfully. {len(files)} file(s) received.",
            "file_name": files[0].filename if files else "homework.pdf",
            "file_url": f"/api/homework/download/{homework_id}",  # Placeholder
            "uploaded_at": datetime.now().isoformat()
        }
        
    except Exception as e:
        import traceback
        error_msg = f"Error processing homework submission: {str(e)}\n\nTraceback:\n{traceback.format_exc()}"
        print(error_msg)
        return {
            "error": f"Error processing homework: {str(e)}",
            "homework_id": f"hw_error_{int(time.time())}",
            "subject": subject if 'subject' in locals() else "Unknown",
            "status": "error",
            "message": "An error occurred while uploading your homework. Please try again.",
            "file_name": "error.pdf",
            "file_url": ""
        }

@app.post("/api/homework/generate")
async def generate_homework_document(req: HomeworkAssignmentRequest):
    """Generate a Word document homework assignment with multiple questions and marks"""
    if not DOCX_AVAILABLE:
        return {"error": "python-docx library not available. Please install: pip install python-docx"}
    
    subject = req.subject or "Mathematics"
    weak_areas = req.weak_areas or []
    num_questions = max(3, min(req.num_questions, 15))  # Between 3-15 questions
    
    # Generate questions using AI based on weak areas
    questions = []
    total_marks = 0
    
    # If weak areas provided, generate questions for each
    if weak_areas:
        questions_per_area = max(1, num_questions // len(weak_areas))
        for area in weak_areas[:5]:  # Max 5 weak areas
            area_questions = generate_questions_for_topic(subject, area, questions_per_area)
            questions.extend(area_questions)
    else:
        # Generate general questions
        questions = generate_questions_for_topic(subject, "general practice", num_questions)
    
    # Limit to requested number
    questions = questions[:num_questions]
    
    # Calculate total marks
    total_marks = sum(q.get('marks', 10) for q in questions)
    
    # Create Word document
    doc = Document()
    
    # Title
    title = doc.add_heading(f'{subject} Homework Assignment', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header information
    doc.add_paragraph(f'Subject: {subject}')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph(f'Total Marks: {total_marks}')
    doc.add_paragraph(f'Deadline: {(datetime.now().replace(day=datetime.now().day+7)).strftime("%B %d, %Y")}')
    doc.add_paragraph('')
    
    # Instructions
    doc.add_heading('Instructions:', level=1)
    instructions = doc.add_paragraph()
    instructions.add_run('1. Answer all questions clearly and show all working.')
    doc.add_paragraph('2. Submit your completed homework through the Homework section of the platform.')
    doc.add_paragraph('3. Ensure your solutions are neat and well-organized.')
    doc.add_paragraph('')
    
    # Questions
    doc.add_heading('Questions:', level=1)
    
    for i, q in enumerate(questions, 1):
        # Question number and text
        q_para = doc.add_paragraph()
        q_para.add_run(f'Question {i}').bold = True
        q_para.add_run(f' ({q.get("marks", 10)} marks)')
        doc.add_paragraph(q['question'])
        doc.add_paragraph('')  # Space for answer
        
        # Add sub-questions if any
        if 'sub_questions' in q:
            for j, sq in enumerate(q['sub_questions'], 1):
                sq_para = doc.add_paragraph()
                sq_para.add_run(f'  {i}.{j}').bold = True
                sq_para.add_run(f' ({sq.get("marks", 5)} marks)')
                doc.add_paragraph(f'  {sq["question"]}')
                doc.add_paragraph('')
    
    # Footer
    doc.add_paragraph('')
    doc.add_paragraph('Good luck!', style='Intense Quote')
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    # Generate filename
    filename = f"{subject}_Homework_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    # Return file
    return FileResponse(
        temp_file.name,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=filename,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

def generate_questions_for_topic(subject: str, topic: str, num_questions: int) -> List[dict]:
    """Generate questions for a specific topic using AI"""
    questions = []
    
    # Subject-specific question templates
    templates = {
        "Mathematics": {
            "algebra": [
                {"question": "Solve the equation: 2x + 5 = 15. Show all steps.", "marks": 5},
                {"question": "Factorize: x² + 7x + 12", "marks": 6},
                {"question": "Solve the simultaneous equations: 2x + 3y = 12 and x - y = 1", "marks": 8}
            ],
            "calculus": [
                {"question": "Find the derivative of f(x) = x³ + 2x² - 5x + 3", "marks": 6},
                {"question": "Evaluate the integral: ∫(2x + 3)dx", "marks": 7},
                {"question": "Find the gradient of the curve y = x² at the point (2, 4)", "marks": 6}
            ],
            "geometry": [
                {"question": "Calculate the area of a circle with radius 7cm. (Use π = 22/7)", "marks": 5},
                {"question": "A triangle has sides of length 5cm, 12cm, and 13cm. Determine if it is a right-angled triangle.", "marks": 7},
                {"question": "Find the volume of a cylinder with radius 4cm and height 10cm.", "marks": 6}
            ]
        },
        "Accounting": {
            "double_entry": [
                {"question": "Record the following transaction using double-entry bookkeeping: Purchased goods worth $500 on credit from ABC Suppliers.", "marks": 8},
                {"question": "Journalize: Paid rent of $200 by cash.", "marks": 6},
                {"question": "Record: Sold goods worth $300 for cash.", "marks": 7}
            ],
            "financial_statements": [
                {"question": "Prepare a simple Income Statement from the following: Sales $10,000, Cost of Sales $6,000, Expenses $2,000.", "marks": 10},
                {"question": "Calculate the Gross Profit if Sales = $8,000 and Cost of Sales = $5,000.", "marks": 5}
            ]
        },
        "Economics": {
            "demand_supply": [
                {"question": "Explain what happens to the demand curve when consumer income increases.", "marks": 8},
                {"question": "If the price of a good increases from $10 to $15 and quantity demanded decreases from 100 to 80, calculate the price elasticity of demand.", "marks": 10}
            ],
            "market_structures": [
                {"question": "Compare and contrast perfect competition and monopoly market structures.", "marks": 12}
            ]
        },
        "Computer Science": {
            "algorithms": [
                {"question": "Write an algorithm to find the largest number in a list of 10 numbers.", "marks": 8},
                {"question": "Explain the time complexity of binary search algorithm.", "marks": 7}
            ],
            "programming": [
                {"question": "Write a Python function to calculate the factorial of a number.", "marks": 10}
            ]
        },
        "Business Studies": {
            "swot": [
                {"question": "Conduct a SWOT analysis for a local retail business.", "marks": 12},
                {"question": "Explain the 4 Ps of marketing mix with examples.", "marks": 10}
            ],
            "management": [
                {"question": "Compare autocratic and democratic leadership styles.", "marks": 9}
            ]
        }
    }
    
    # Get templates for subject
    subject_templates = templates.get(subject, templates["Mathematics"])
    
    # Select questions based on topic
    if topic.lower() in subject_templates:
        topic_questions = subject_templates[topic.lower()]
    else:
        # Use first available topic
        topic_questions = list(subject_templates.values())[0] if subject_templates else []
    
    # Generate additional questions using AI if needed
    if len(topic_questions) < num_questions:
        # Use AI to generate more questions
        prompt = f"""Generate {num_questions - len(topic_questions)} additional {subject} homework questions on the topic: {topic}.
        
        Each question should:
        - Be appropriate for Forms 5-6 level students in Zimbabwe
        - Include allocated marks (between 5-15 marks per question)
        - Be clear and specific
        - Test understanding of the topic
        
        Format as JSON array with structure:
        [{{"question": "question text", "marks": 10}}]
        
        Return only the JSON array, no other text."""
        
        try:
            ai_response = call_ai_with_fallback(prompt, "", subject)
            # Try to parse JSON from response
            import re
            json_match = re.search(r'\[.*\]', ai_response, re.DOTALL)
            if json_match:
                ai_questions = json.loads(json_match.group())
                topic_questions.extend(ai_questions)
        except:
            pass
    
    # Return requested number of questions
    return topic_questions[:num_questions]

@app.get("/")
def root():
    return {"message": "AI Learning Platform API", "status": "running"}

@app.get("/health")
def health():
    return {"status": "healthy"}

class QuizQuestionRequest(BaseModel):
    subject: str
    num_questions: int = 20
    set_number: int = 1
    topic: Optional[str] = None

class QuizQuestionResponse(BaseModel):
    questions: List[dict]
    set_number: int
    total_sets: int
    subject: str

@app.post("/api/quiz/generate-questions", response_model=QuizQuestionResponse)
async def generate_quiz_questions(req: QuizQuestionRequest):
    """Generate quiz questions from syllabus content using AI"""
    subject = req.subject or "Mathematics"
    num_questions = max(20, min(req.num_questions, 50))  # Between 20-50 questions per set
    set_number = req.set_number or 1
    
    # Get syllabus content for context
    chunks = get_chunks(subject)
    context = ""
    if chunks:
        # Use multiple chunks for better question diversity
        context = "\n\n".join(chunks[:5])  # Use top 5 chunks
    
    # Build prompt for question generation
    topic_context = f" on the topic: {req.topic}" if req.topic else ""
    prompt = f"""You are an expert quiz creator for Forms 5-6 {subject} students in Zimbabwe.

**Task**: Generate {num_questions} multiple-choice quiz questions{topic_context} based on the syllabus content provided below.

**Syllabus Content**:
{context[:3000] if context else "General {subject} curriculum for Forms 5-6"}

**Requirements**:
1. Each question must have exactly 4 options (A, B, C, D)
2. Only ONE option should be correct
3. Questions should test understanding, not just memorization
4. Difficulty should be appropriate for Forms 5-6 level
5. Cover different topics from the syllabus
6. Make questions clear and unambiguous
7. Options should be plausible (not obviously wrong)

**Output Format** (JSON array):
[
  {{
    "question": "Question text here?",
    "options": ["Option A", "Option B", "Option C", "Option D"],
    "correct": 0,
    "explanation": "Brief explanation of the correct answer"
  }},
  ...
]

Return ONLY the JSON array, no other text."""

    # Generate questions using AI
    try:
        ai_response = call_ai_with_fallback(prompt, context, subject)
        
        # Try to parse JSON from response
        import re
        json_match = re.search(r'\[.*\]', ai_response, re.DOTALL)
        if json_match:
            questions = json.loads(json_match.group())
            # Ensure we have the right number
            questions = questions[:num_questions]
            
            # If we got fewer questions, generate more
            while len(questions) < num_questions:
                # Generate additional questions
                additional_prompt = f"""Generate {num_questions - len(questions)} more {subject} quiz questions. Return as JSON array with same format as before."""
                additional_response = call_ai_with_fallback(additional_prompt, context, subject)
                additional_match = re.search(r'\[.*\]', additional_response, re.DOTALL)
                if additional_match:
                    additional_questions = json.loads(additional_match.group())
                    questions.extend(additional_questions[:num_questions - len(questions)])
                else:
                    break
            
            # Calculate total sets (estimate: assume we can generate many sets from syllabus)
            # For now, return a reasonable number
            total_sets = max(5, len(chunks) // 4) if chunks else 5
            
            return QuizQuestionResponse(
                questions=questions[:num_questions],
                set_number=set_number,
                total_sets=total_sets,
                subject=subject
            )
    except Exception as e:
        print(f"Error generating questions: {e}")
        import traceback
        traceback.print_exc()
    
    # Fallback: return empty questions
    return QuizQuestionResponse(
        questions=[],
        set_number=set_number,
        total_sets=5,
        subject=subject
    )

class DocumentRequest(BaseModel):
    subject: str
    answer: str
    question: str = ""

@app.post("/api/chat/download-document")
async def download_chat_document(req: DocumentRequest):
    """Generate and download a Word document from chat response"""
    if not DOCX_AVAILABLE:
        return {"error": "python-docx library not available"}
    
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'{req.subject} - Solution', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').bold = True
        if req.question:
            doc.add_paragraph(f'Question: {req.question}')
        doc.add_paragraph('')
        
        # Format content based on subject
        content = format_answer_for_word_document(req.answer, req.subject)
        
        # Split content into paragraphs and format
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Check if it's a heading
            if 'INCOME STATEMENT' in para.upper() or 'Statement of Financial Position' in para:
                heading = doc.add_heading(para, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif para.startswith('Step') or para.startswith('Question') or para.startswith('Solution'):
                doc.add_heading(para, level=2)
            elif para.startswith('Explanation:'):
                p = doc.add_paragraph(para)
                p.style = 'Intense Quote'
            else:
                doc.add_paragraph(para)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_path = temp_file.name
        temp_file.close()
        doc.save(temp_path)
        
        # Return file
        from fastapi.background import BackgroundTask
        def cleanup(temp_path):
            if os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except:
                    pass
        
        return FileResponse(
            temp_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f'{req.subject}_Solution_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            background=BackgroundTask(lambda: os.unlink(temp_path) if os.path.exists(temp_path) else None)
        )
    except Exception as e:
        return {"error": f"Error generating document: {str(e)}"}

@app.post("/api/chat/download-pdf")
async def download_chat_pdf(req: DocumentRequest):
    """Generate and download a PDF document from chat response"""
    try:
        from reportlab.lib.pagesizes import letter, A4  # type: ignore[reportMissingModuleSource]
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore[reportMissingModuleSource]
        from reportlab.lib.units import inch  # type: ignore[reportMissingModuleSource]
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak  # type: ignore[reportMissingModuleSource]
        from reportlab.lib.enums import TA_LEFT, TA_CENTER  # type: ignore[reportMissingModuleSource]
        from reportlab.pdfbase import pdfmetrics  # type: ignore[reportMissingModuleSource]
        from reportlab.pdfbase.ttfonts import TTFont  # type: ignore[reportMissingModuleSource]
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_path = temp_file.name
        temp_file.close()
        
        # Create PDF document
        doc = SimpleDocTemplate(temp_path, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)
        
        # Container for the 'Flowable' objects
        story = []
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#667eea',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        # Add title
        title = Paragraph(f'{req.subject} - Solution', title_style)
        story.append(title)
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata
        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontSize=10,
            textColor='#666'
        )
        story.append(Paragraph(f'<b>Generated:</b> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', meta_style))
        if req.question:
            story.append(Paragraph(f'<b>Question:</b> {req.question}', meta_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Format content
        content = format_answer_for_word_document(req.answer, req.subject)
        
        # Split into paragraphs and add to story
        paragraphs = content.split('\n\n')
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=12
        )
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Escape HTML entities
            para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Check for headings
            if para.upper().startswith('INCOME STATEMENT') or 'Statement of Financial Position' in para:
                heading_style = ParagraphStyle(
                    'SectionHeading',
                    parent=styles['Heading2'],
                    fontSize=14,
                    textColor='#333',
                    spaceAfter=12,
                    alignment=TA_CENTER
                )
                story.append(Paragraph(para, heading_style))
            elif para.startswith('Step') or para.startswith('Question') or para.startswith('Solution'):
                heading_style = ParagraphStyle(
                    'SubHeading',
                    parent=styles['Heading3'],
                    fontSize=12,
                    textColor='#667eea',
                    spaceAfter=8
                )
                story.append(Paragraph(para, heading_style))
            else:
                story.append(Paragraph(para, body_style))
            
            story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        
        # Return file
        from fastapi.background import BackgroundTask
        return FileResponse(
            temp_path,
            media_type='application/pdf',
            filename=f'{req.subject}_Solution_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
            background=BackgroundTask(lambda: os.unlink(temp_path) if os.path.exists(temp_path) else None)
        )
    except ImportError:
        return {"error": "reportlab library not available. Install with: pip install reportlab"}
    except Exception as e:
        import traceback
        return {"error": f"Error generating PDF: {str(e)}\n{traceback.format_exc()}"}

def format_answer_for_word_document(answer: str, subject: str) -> str:
    """Format answer text for Word document (plain text, no HTML)"""
    # Remove HTML tags but preserve structure
    import html
    text = html.unescape(answer)
    text = re.sub(r'<[^>]+>', '', text)  # Remove HTML tags
    
    # Clean up formatting
    text = re.sub(r'\n{3,}', '\n\n', text)  # Remove excessive newlines
    
    return text

